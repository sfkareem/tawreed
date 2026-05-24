import os
import json
import re
import uuid
import logging
import datetime
import base64
import requests
import traceback
from typing import Dict, List, Any, Tuple
import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
import pdfplumber
import docx
import fitz  # PyMuPDF

# Constants and Configuration Setup
CONFIG_DIR = os.path.expanduser(r"~\.tawreed")
CONFIG_FILE = os.path.join(CONFIG_DIR, "config.json")
LOG_DIR = os.path.join(CONFIG_DIR, "logs")

# Create directories if they do not exist
os.makedirs(CONFIG_DIR, exist_ok=True)
os.makedirs(LOG_DIR, exist_ok=True)

# Setup Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(os.path.join(LOG_DIR, "tawreed.log"), encoding="utf-8"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("TawreedBackend")

DEFAULT_CONFIG = {
    "api_provider": "gemini",
    "base_url": "https://generativelanguage.googleapis.com",
    "api_key": "",
    "model_name": "gemini-1.5-flash",
    "preferred_language": "bilingual",
    "max_file_size_mb": 15,
    "max_content_length_chars": 80000
}

def load_config() -> Dict[str, Any]:
    if not os.path.exists(CONFIG_FILE):
        save_config(DEFAULT_CONFIG)
        return DEFAULT_CONFIG
    try:
        with open(CONFIG_FILE, "r", encoding="utf-8") as f:
            config = json.load(f)
            # Ensure all keys are present
            for k, v in DEFAULT_CONFIG.items():
                if k not in config:
                    config[k] = v
            return config
    except Exception as e:
        logger.error(f"Error loading config: {e}")
        return DEFAULT_CONFIG

def save_config(config: Dict[str, Any]) -> bool:
    try:
        with open(CONFIG_FILE, "w", encoding="utf-8") as f:
            json.dump(config, f, indent=4, ensure_ascii=False)
        return True
    except Exception as e:
        logger.error(f"Error saving config: {e}")
        return False

# File Parsers and Serializers
class DocumentParser:
    @staticmethod
    def parse_excel(file_path: str) -> Tuple[str, List[str]]:
        """Parses all sheets of an Excel file and serializes to text."""
        warnings = []
        serialized = []
        try:
            xls = pd.ExcelFile(file_path)
            serialized.append(f"Source Excel: {os.path.basename(file_path)}\n")
            for sheet_name in xls.sheet_names:
                df = pd.read_excel(xls, sheet_name=sheet_name)
                serialized.append(f"--- Sheet: {sheet_name} ---")
                
                # Check empty
                if df.empty:
                    serialized.append("[Empty Sheet]")
                    continue
                
                # Drop rows/columns that are completely empty
                df = df.dropna(how='all').dropna(axis=1, how='all')
                
                # Convert to readable rows
                # Include column headers
                headers = [str(col) for col in df.columns]
                serialized.append("Columns: | " + " | ".join(headers) + " |")
                
                for idx, row in df.iterrows():
                    row_vals = []
                    for val in row:
                        if pd.isna(val):
                            row_vals.append("")
                        else:
                            # Format floats nicely
                            if isinstance(val, float):
                                if val.is_integer():
                                    row_vals.append(str(int(val)))
                                else:
                                    row_vals.append(f"{val:.2f}")
                            else:
                                row_vals.append(str(val))
                    # Row coordinates start at index+2 usually (accounting for headers and 0-indexing)
                    serialized.append(f"Row {idx+2}: | " + " | ".join(row_vals) + " |")
                
                serialized.append("") # Empty line between sheets
        except Exception as e:
            logger.error(f"Excel parsing error: {e}")
            warnings.append(f"Excel read warning: {str(e)}")
            serialized.append(f"[Error reading Excel file: {str(e)}]")
            
        return "\n".join(serialized), warnings

    @staticmethod
    def parse_csv(file_path: str) -> Tuple[str, List[str]]:
        warnings = []
        serialized = []
        try:
            df = pd.read_csv(file_path)
            serialized.append(f"Source CSV: {os.path.basename(file_path)}\n")
            df = df.dropna(how='all')
            headers = [str(col) for col in df.columns]
            serialized.append("Columns: | " + " | ".join(headers) + " |")
            for idx, row in df.iterrows():
                row_vals = ["" if pd.isna(v) else str(v) for v in row]
                serialized.append(f"Row {idx+2}: | " + " | ".join(row_vals) + " |")
        except Exception as e:
            logger.error(f"CSV parsing error: {e}")
            warnings.append(f"CSV read warning: {str(e)}")
            serialized.append(f"[Error reading CSV file: {str(e)}]")
        return "\n".join(serialized), warnings

    @staticmethod
    def parse_word(file_path: str) -> Tuple[str, List[str]]:
        warnings = []
        serialized = []
        try:
            doc = docx.Document(file_path)
            serialized.append(f"Source Word Document: {os.path.basename(file_path)}\n")
            
            # Read paragraphs and tables in order of appearance
            # We will process paragraphs first, then tables.
            serialized.append("=== Paragraphs ===")
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    serialized.append(f"P{i+1}: {text}")
            
            serialized.append("\n=== Tables ===")
            for t_idx, table in enumerate(doc.tables):
                serialized.append(f"Table {t_idx+1}:")
                for r_idx, row in enumerate(table.rows):
                    row_vals = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    serialized.append(f"Row {r_idx+1}: | " + " | ".join(row_vals) + " |")
                serialized.append("")
        except Exception as e:
            logger.error(f"Word parsing error: {e}")
            warnings.append(f"Word read warning: {str(e)}")
            serialized.append(f"[Error reading Word file: {str(e)}]")
        return "\n".join(serialized), warnings

    @staticmethod
    def parse_text(file_path: str) -> Tuple[str, List[str]]:
        warnings = []
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return f"Source Text File: {os.path.basename(file_path)}\n\n{content}", warnings
        except Exception as e:
            logger.error(f"Text parsing error: {e}")
            warnings.append(f"Text read warning: {str(e)}")
            return f"[Error reading text file: {str(e)}]", warnings

    @staticmethod
    def parse_digital_pdf(file_path: str) -> Tuple[str, List[str], bool]:
        """Extracts text from a digital PDF. Returns content, warnings, and is_scanned flag."""
        warnings = []
        serialized = []
        total_text_length = 0
        try:
            with pdfplumber.open(file_path) as pdf:
                serialized.append(f"Source PDF: {os.path.basename(file_path)}\n")
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text()
                    if text:
                        text_clean = text.strip()
                        total_text_length += len(text_clean)
                        serialized.append(f"--- Page {i+1} ---")
                        serialized.append(text_clean)
                        
                        # Also attempt to parse tables on the page
                        tables = page.extract_tables()
                        if tables:
                            serialized.append("\n[Tables Detected on Page]")
                            for t_idx, table in enumerate(tables):
                                serialized.append(f"Table {t_idx+1}:")
                                for r_idx, row in enumerate(table):
                                    row_vals = ["" if v is None else str(v).strip().replace("\n", " ") for v in row]
                                    serialized.append(f"Row {r_idx+1}: | " + " | ".join(row_vals) + " |")
                                serialized.append("")
                    else:
                        serialized.append(f"--- Page {i+1} (No text extracted) ---")
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            warnings.append(f"PDF read warning: {str(e)}")
            return f"[Error reading PDF: {str(e)}]", warnings, True
        
        # If very little text is extracted relative to page count, it's likely scanned
        is_scanned = total_text_length < 100 * len(pdf.pages) if 'pdf' in locals() else True
        return "\n".join(serialized), warnings, is_scanned

    @staticmethod
    def render_pdf_pages_to_images(file_path: str) -> List[Tuple[int, bytes]]:
        """Renders PDF pages as JPEG image bytes using PyMuPDF (fitz) for AI Vision OCR."""
        image_pages = []
        try:
            doc = fitz.open(file_path)
            for i, page in enumerate(doc):
                # Render page at 150 DPI (good balance of quality and size)
                pix = page.get_pixmap(dpi=150)
                img_bytes = pix.tobytes("jpeg")
                image_pages.append((i + 1, img_bytes))
            doc.close()
        except Exception as e:
            logger.error(f"Error rendering PDF pages to images: {e}")
        return image_pages

# AI Invoker Service
class AIService:
    @staticmethod
    def get_system_prompt(language: str) -> str:
        import sys
        
        # Load SKILL.md dynamically
        skill_content = ""
        try:
            # Try same folder as tawreed_backend.py
            base_dir = os.path.dirname(os.path.abspath(__file__))
            skill_path = os.path.join(base_dir, "SKILL.md")
            if not os.path.exists(skill_path):
                # Try sys._MEIPASS if bundled by PyInstaller
                meipass = getattr(sys, '_MEIPASS', None)
                if meipass:
                    skill_path = os.path.join(meipass, "SKILL.md")
            if not os.path.exists(skill_path):
                # Try current working directory
                skill_path = "SKILL.md"
                
            if os.path.exists(skill_path):
                with open(skill_path, "r", encoding="utf-8") as sf:
                    skill_content = sf.read()
            else:
                logger.warning(f"SKILL.md not found at path: {os.path.abspath(skill_path)}")
        except Exception as se:
            logger.warning(f"Could not load SKILL.md dynamically: {se}")

        lang_instruction = ""
        schema_fields = ""
        
        if language == "arabic":
            lang_instruction = "The final output materials name, specification, quantity basis, notes, and assumptions MUST be extracted or translated into Arabic (العربية)."
            schema_fields = """
            - package: Material package/category (e.g. أعمال البلاط، الأبواب، أعمال الجبس)
            - material_name: Specific material name in Arabic
            - specification: Detailed specifications, sizes, standards, grades, or accessories in Arabic
            - unit: Unit of measurement (e.g. م2، م.ط، عدد، طن)
            - estimated_quantity: Numeric quantity, or null if not found
            - quantity_basis: Where quantity came from or how calculated (in Arabic)
            - related_boq_items: BOQ item numbers or codes
            - source_sections: Section or sheet name in the document
            - confidence: Confidence level (High, Medium, Low)
            - notes: General notes or missing accessories (in Arabic)
            """
        elif language == "english":
            lang_instruction = "The final output materials name, specification, quantity basis, notes, and assumptions MUST be extracted or translated into English."
            schema_fields = """
            - package: Material package/category (e.g. Flooring, Doors, Ceiling, Painting)
            - material_name: Specific material name in English
            - specification: Detailed specifications, sizes, standards, grades, or accessories in English
            - unit: Unit of measurement (e.g. m2, m, pcs, ton)
            - estimated_quantity: Numeric quantity, or null if not found
            - quantity_basis: Where quantity came from or how calculated (in English)
            - related_boq_items: BOQ item numbers or codes
            - source_sections: Section or sheet name in the document
            - confidence: Confidence level (High, Medium, Low)
            - notes: General notes or missing accessories (in English)
            """
        else: # bilingual
            lang_instruction = "The final output MUST contain separate bilingual fields for English and Arabic. Do not merge them in a single column; write them in the respective fields."
            schema_fields = """
            - package_en: Material package/category in English (e.g. Flooring, Doors)
            - package_ar: Material package/category in Arabic (e.g. أعمال البلاط، الأبواب)
            - material_name_en: Material name in English
            - material_name_ar: Material name in Arabic
            - specification_en: Specifications/details in English
            - specification_ar: Specifications/details in Arabic
            - unit_en: Unit of measurement in English (e.g. m2, pcs)
            - unit_ar: Unit of measurement in Arabic (e.g. م2، حبة)
            - estimated_quantity: Numeric quantity, or null if not found
            - quantity_basis: Where quantity came from or how calculated (bilingual or English)
            - related_boq_items: BOQ item codes or numbers
            - source_sections: Section or sheet name in the document
            - confidence: Confidence level (High, Medium, Low)
            - notes: General notes or missing accessories (bilingual)
            """

        prompt = f"""You are a professional Quantity Surveyor, Estimator, and Construction Procurement Engineer.
Your job is to analyze the provided project document (BOQ, Quotation, cost breakdown, or drawings pages) and extract a structured Materials List for procurement.

Rules for Extraction:
1. Identify all materials required for the project. A single BOQ description may contain multiple materials (e.g., a door item might include the wood door, door frame, ironmongery, and painting). Break them down into separate rows if they are distinct procurement items.
2. Deduplicate: If the same material with the exact same specification appears in multiple items, group them together, but list all related BOQ items in the 'related_boq_items' column.
3. Clean specs: Do not copy the entire BOQ scope (like testing, commissioning, protection, and installation) into the material specification. Only extract the material attributes (dimensions, grade, thickness, finish, standards, etc.).
4. Units: Extract standard procurement units.
5. Quantity Basis: Clearly state if the quantity is copied directly from the BOQ, calculated, or assumed.
6. Package/Category: Categorize each material under standard packages (e.g., Concrete, Masonry, Waterproofing, Plastering, Painting, Flooring, Gypsum Ceiling, Joinery, Metal Works, Sanitary Ware, Glazing, etc.).
7. Flags & Assumptions: If details like thickness, grade, or quantities are missing, make a logical engineering assumption, record it in the assumptions list, and flag the material as 'needs manual review' in the review flags.

Language Directive:
{lang_instruction}

Output Schema:
You MUST respond with a single valid JSON object containing exactly four lists:
{{
  "materials": [
    // Array of objects with the following fields:
    {schema_fields}
  ],
  "summary": {{
    "total_materials_extracted": 0,
    "packages_detected": ["Package A", "Package B"],
    "high_confidence_count": 0,
    "medium_confidence_count": 0,
    "low_confidence_count": 0,
    "review_flags_count": 0
  }},
  "review_flags": [
    // Array of objects for items requiring critical engineer review:
    {{
      "material_name": "Name of the flagged material",
      "issue_type": "e.g. Missing Specification, Unclear Quantity, Potential Duplicate, Missing Accessories",
      "description": "Explanation of what is missing or needs checking",
      "severity": "High or Medium"
    }}
  ],
  "assumptions": [
    // Array of objects for assumptions made during parsing:
    {{
      "scope_or_material": "Item name or category",
      "assumption_made": "What was assumed (e.g. assumed thickness of glass is 12mm as not specified)",
      "justification": "Why this assumption makes sense"
    }}
  ],
  "warnings": [
    // Array of strings representing data warnings:
    "e.g. Sheet X was skipped because it was empty",
    "e.g. Quantity in row 12 could not be parsed as a number"
  ]
}}

CRITICAL: Return ONLY valid, minified JSON. Do not write introductory text, markdown formatting codeblocks, or trailing explanations. Start with '{{' and end with '}}'."""
        
        if skill_content:
            prompt += f"\n\nAdditionally, you must strictly follow these spreadsheet formatting and styling instructions:\n\n{skill_content}"
            
        return prompt

    @staticmethod
    def call_ai(config: Dict[str, Any], prompt: str, text_content: str = "", image_bytes_list: List[bytes] = None) -> Tuple[str, str]:
        """Calls the configured AI API. Returns (raw_response, error_message)."""
        provider = config.get("api_provider", "gemini")
        api_key = config.get("api_key", "")
        base_url = config.get("base_url", "")
        model = config.get("model_name", "")
        
        if not api_key:
            return "", "API Key is missing. Please set it in Settings."

        try:
            if provider == "gemini":
                # Google Gemini API
                # By default, Gemini base url is https://generativelanguage.googleapis.com
                url = f"{base_url}/v1beta/models/{model}:generateContent?key={api_key}"
                
                # Build contents array
                parts = [{"text": prompt}]
                if text_content:
                    parts.append({"text": f"Here is the document content to analyze:\n\n{text_content}"})
                
                if image_bytes_list:
                    for img_bytes in image_bytes_list:
                        encoded = base64.b64encode(img_bytes).decode("utf-8")
                        parts.append({
                            "inlineData": {
                                "mimeType": "image/jpeg",
                                "data": encoded
                            }
                        })
                
                payload = {
                    "contents": [{
                        "parts": parts
                    }],
                    "generationConfig": {
                        "responseMimeType": "application/json"
                    }
                }
                
                logger.info(f"Sending request to Gemini API (Model: {model})...")
                res = requests.post(url, json=payload, headers={"Content-Type": "application/json"}, timeout=120)
                
                if res.status_code != 200:
                    logger.error(f"Gemini API returned error {res.status_code}: {res.text}")
                    return "", f"API Error {res.status_code}: {res.text}"
                
                data = res.json()
                try:
                    text_out = data['candidates'][0]['content']['parts'][0]['text']
                    return text_out, ""
                except Exception as parse_e:
                    logger.error(f"Failed to parse Gemini response payload: {data}. Error: {parse_e}")
                    return "", "Invalid response format from Gemini API."

            elif provider == "openai":
                # OpenAI API
                url = f"{base_url}/v1/chat/completions" if base_url else "https://api.openai.com/v1/chat/completions"
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                }
                
                messages = [
                    {"role": "system", "content": prompt}
                ]
                
                # Construct user content (which can be multimodal)
                user_content = []
                if text_content:
                    user_content.append({"type": "text", "text": f"Here is the document content to analyze:\n\n{text_content}"})
                
                if image_bytes_list:
                    for img_bytes in image_bytes_list:
                        encoded = base64.b64encode(img_bytes).decode("utf-8")
                        user_content.append({
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{encoded}"
                            }
                        })
                
                messages.append({"role": "user", "content": user_content})
                
                payload = {
                    "model": model,
                    "messages": messages,
                    "response_format": {"type": "json_object"}
                }
                
                logger.info(f"Sending request to OpenAI API (Model: {model})...")
                res = requests.post(url, json=payload, headers=headers, timeout=120)
                
                if res.status_code != 200:
                    logger.error(f"OpenAI API returned error {res.status_code}: {res.text}")
                    return "", f"API Error {res.status_code}: {res.text}"
                
                data = res.json()
                try:
                    text_out = data['choices'][0]['message']['content']
                    return text_out, ""
                except Exception as parse_e:
                    logger.error(f"Failed to parse OpenAI response payload: {data}. Error: {parse_e}")
                    return "", "Invalid response format from OpenAI API."

            elif provider == "anthropic":
                # Anthropic Claude API
                # Claude uses base_url = https://api.anthropic.com
                url = f"{base_url}/v1/messages" if base_url else "https://api.anthropic.com/v1/messages"
                headers = {
                    "x-api-key": api_key,
                    "anthropic-version": "2023-06-01",
                    "Content-Type": "application/json"
                }
                
                # Construct content list
                content_list = []
                if text_content:
                    content_list.append({"type": "text", "text": f"Here is the document content to analyze:\n\n{text_content}"})
                
                if image_bytes_list:
                    for img_bytes in image_bytes_list:
                        encoded = base64.b64encode(img_bytes).decode("utf-8")
                        content_list.append({
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": "image/jpeg",
                                "data": encoded
                            }
                        })
                
                payload = {
                    "model": model,
                    "system": prompt,
                    "messages": [{"role": "user", "content": content_list}],
                    "max_tokens": 4096
                }
                
                logger.info(f"Sending request to Anthropic API (Model: {model})...")
                res = requests.post(url, json=payload, headers=headers, timeout=120)
                
                if res.status_code != 200:
                    logger.error(f"Anthropic API returned error {res.status_code}: {res.text}")
                    return "", f"API Error {res.status_code}: {res.text}"
                
                data = res.json()
                try:
                    text_out = data['content'][0]['text']
                    return text_out, ""
                except Exception as parse_e:
                    logger.error(f"Failed to parse Anthropic response payload: {data}. Error: {parse_e}")
                    return "", "Invalid response format from Anthropic API."

            else: # custom OpenAI-compatible endpoint
                url = f"{base_url}/chat/completions" if base_url else "http://localhost:11434/v1/chat/completions"
                headers = {
                    "Content-Type": "application/json"
                }
                if api_key:
                    headers["Authorization"] = f"Bearer {api_key}"
                
                # Custom endpoints might not support Vision. If image exists and model is custom,
                # we can warn or send. Let's send anyway.
                user_content = []
                if text_content:
                    user_content.append({"type": "text", "text": f"Here is the document content to analyze:\n\n{text_content}"})
                if image_bytes_list:
                    for img_bytes in image_bytes_list:
                        encoded = base64.b64encode(img_bytes).decode("utf-8")
                        user_content.append({
                            "type": "image_url",
                            "image_url": {
                                "url": f"data:image/jpeg;base64,{encoded}"
                            }
                        })
                
                payload = {
                    "model": model,
                    "messages": [
                        {"role": "system", "content": prompt},
                        {"role": "user", "content": user_content if len(user_content) > 1 or image_bytes_list else (text_content or "")}
                    ]
                }
                
                logger.info(f"Sending request to Custom API (Model: {model}) at URL: {url}...")
                res = requests.post(url, json=payload, headers=headers, timeout=120)
                
                if res.status_code != 200:
                    logger.error(f"Custom API returned error {res.status_code}: {res.text}")
                    return "", f"API Error {res.status_code}: {res.text}"
                
                data = res.json()
                try:
                    text_out = data['choices'][0]['message']['content']
                    return text_out, ""
                except Exception as parse_e:
                    logger.error(f"Failed to parse Custom API response payload: {data}. Error: {parse_e}")
                    return "", "Invalid response format from Custom API."
                    
        except Exception as e:
            logger.error(f"HTTP call to AI failed: {e}\n{traceback.format_exc()}")
            return "", f"Network/Connection Error: {str(e)}"

    @staticmethod
    def get_print_settings(config: Dict[str, Any], columns: List[str], data_sample: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Queries the LLM for print styling configurations to optimize PDF layout."""
        fallback = {
            "orientation": "landscape" if len(columns) > 6 else "portrait",
            "font_size": "11px",
            "column_widths": {col: f"{int(100/len(columns))}%" for col in columns},
            "wrap_columns": [col for col in columns if "spec" in col or "note" in col or "basis" in col]
        }
        
        # If API key is missing, return fallback immediately
        if not config.get("api_key", ""):
            return fallback
            
        prompt = f"""You are a print layout designer and spreadsheet formatting expert.
We are printing a materials quantity takeoff spreadsheet to a PDF report.
To make it fit nicely, please analyze the columns and the sample data below to choose the best print settings:

Columns: {columns}
Sample Data:
{json.dumps(data_sample[:5], indent=2, ensure_ascii=False)}

Choose:
1. "orientation": "landscape" or "portrait". Use landscape if there are many columns (e.g., bilingual or > 6 columns) or if the specification column requires significant horizontal space to prevent excessive line wrapping.
2. "font_size": a CSS font size like "10px", "11px", or "12px".
3. "column_widths": a JSON dictionary mapping the column keys to recommended percentage widths (e.g., {{"material_name_en": "20%", "specification_en": "35%"}}). The sum of percentages should equal 100%.
4. "wrap_columns": a list of column keys that have long text and MUST be wrapped (e.g., ["specification_en", "notes"]).

Return ONLY a valid JSON object matching this schema, start with '{{' and end with '}}'. Do not write introductory text, markdown formatting codeblocks, or trailing explanations:
{{
  "orientation": "landscape",
  "font_size": "10px",
  "column_widths": {{
     "col_key_1": "15%",
     "col_key_2": "30%"
  }},
  "wrap_columns": ["col_key_2"]
}}
"""
        try:
            logger.info("Calling LLM to optimize print settings...")
            raw_res, err = AIService.call_ai(config, prompt)
            if err:
                logger.warning(f"Print optimizer AI call failed: {err}. Using fallback.")
                return fallback
                
            # Extract JSON block
            json_str = JSONRepairService.extract_json_string(raw_res)
            # Remove trailing commas
            json_str = re.sub(r",\s*(\]|\})", r"\1", json_str)
            settings = json.loads(json_str)
            
            # Enforce validation
            if "orientation" not in settings or settings["orientation"] not in ["landscape", "portrait"]:
                settings["orientation"] = fallback["orientation"]
            if "font_size" not in settings:
                settings["font_size"] = fallback["font_size"]
            if "column_widths" not in settings or not isinstance(settings["column_widths"], dict):
                settings["column_widths"] = fallback["column_widths"]
            if "wrap_columns" not in settings or not isinstance(settings["wrap_columns"], list):
                settings["wrap_columns"] = fallback["wrap_columns"]
                
            logger.info(f"Optimized print settings obtained successfully: {settings}")
            return settings
        except Exception as e:
            logger.warning(f"Failed to query/parse optimized print settings: {e}. Using fallback.")
            return fallback

# JSON Repair Tool
class JSONRepairService:
    @staticmethod
    def extract_json_string(raw_str: str) -> str:
        """Finds and extracts the first matching JSON block from text."""
        raw_str = raw_str.strip()
        # Look for markdown code blocks
        match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", raw_str, re.DOTALL)
        if match:
            return match.group(1).strip()
        
        # Look for first { and last }
        start_idx = raw_str.find("{")
        end_idx = raw_str.rfind("}")
        if start_idx != -1 and end_idx != -1 and end_idx > start_idx:
            return raw_str[start_idx:end_idx + 1].strip()
            
        return raw_str

    @classmethod
    def repair_json(cls, raw_str: str) -> Dict[str, Any]:
        """Tries to clean and parse JSON. Raises exception if failed."""
        cleaned_str = cls.extract_json_string(raw_str)
        
        # Common formatting cleanup:
        # Remove trailing commas inside lists/objects before closing brackets
        cleaned_str = re.sub(r",\s*(\]|\})", r"\1", cleaned_str)
        # Remove double spaces or newlines inside key-value separator
        cleaned_str = re.sub(r'":\s*\n\s*', '": ', cleaned_str)
        
        # Try loading
        return json.loads(cleaned_str)

    @classmethod
    def attempt_ai_repair(cls, config: Dict[str, Any], malformed_str: str, error_msg: str) -> Dict[str, Any]:
        """Sends malformed JSON back to the AI for a single automated repair attempt."""
        logger.info("JSON malformed. Attempting automated AI repair query...")
        repair_prompt = f"""You returned a response that is not valid JSON. 
Error: {error_msg}

Here is the response you returned:
{malformed_str}

Please repair the JSON formatting so it is completely valid according to standard JSON spec. 
Do not omit any extracted materials, assumptions, warnings or review flags from the original response.
Return ONLY the raw valid JSON. Start with '{{' and end with '}}'."""
        
        raw_out, err = AIService.call_ai(config, repair_prompt)
        if err:
            raise Exception(f"AI repair attempt failed due to connection error: {err}")
            
        return cls.repair_json(raw_out)

# Excel Exporter (Tawreed Signature Design)
class ExcelExporter:
    @staticmethod
    def sanitize_sheet_title(name: str) -> str:
        if not name:
            return "General"
        # Remove characters not allowed in Excel sheet names: \ / ? * [ ] :
        for char in r"\/?:*[]":
            name = name.replace(char, "")
        # Truncate to 30 chars
        return name.strip()[:30]

    @staticmethod
    def export(data: Dict[str, Any], output_path: str, language: str) -> None:
        """Generates a styled, formatted Excel workbook with separate sheets for each package."""
        wb = openpyxl.Workbook()
        
        # Colors (Tawreed Signature: Rust Orange + Slate details)
        RUST_HEADER_FILL = PatternFill(start_color="D35400", end_color="D35400", fill_type="solid") # Rust Orange
        WHITE_FONT = Font(name="Segoe UI", size=11, bold=True, color="FFFFFF")
        ZEBRA_FILL = PatternFill(start_color="F9EBEA", end_color="F9EBEA", fill_type="solid") # Warm light tint
        WHITE_FILL = PatternFill(start_color="FFFFFF", end_color="FFFFFF", fill_type="solid")
        
        REGULAR_FONT = Font(name="Segoe UI", size=10)
        BOLD_FONT = Font(name="Segoe UI", size=10, bold=True)
        HEADER_FONT = Font(name="Segoe UI", size=11, bold=True)
        BLUE_FONT = Font(name="Segoe UI", size=10, color="0000FF") # For supplier inputs (Anthropic spec)
        GREEN_FONT = Font(name="Segoe UI", size=10, color="008000") # For worksheet links
        
        BORDER_SIDE = Side(border_style="thin", color="D5D8DC")
        CELL_BORDER = Border(left=BORDER_SIDE, right=BORDER_SIDE, top=BORDER_SIDE, bottom=BORDER_SIDE)
        
        # 1. Summary sheet (First Sheet)
        ws_summary = wb.active
        ws_summary.title = "Summary" if language != "arabic" else "ملخص"
        ws_summary.views.sheetView[0].showGridLines = True
        if language == "arabic":
            ws_summary.views.sheetView[0].rightToLeft = True
            
        # Group materials by package
        materials = data.get("materials", [])
        grouped_materials = {}
        for mat in materials:
            if language == "arabic":
                pkg = mat.get("package") or "عام"
            elif language == "english":
                pkg = mat.get("package") or "General"
            else: # bilingual
                pkg = mat.get("package_en") or mat.get("package") or "General"
            
            pkg = str(pkg).strip()
            if not pkg:
                pkg = "General" if language != "arabic" else "عام"
                
            if pkg not in grouped_materials:
                grouped_materials[pkg] = []
            grouped_materials[pkg].append(mat)
            
        # Determine column structure for packages
        # Column format: (Header, Field, Width, Align, is_input, is_formula, number_format)
        if language == "arabic":
            cols = [
                ("اسم المادة", "material_name", 25, "left", False, False, None),
                ("المواصفات الفنية", "specification", 40, "left", False, False, None),
                ("الوحدة", "unit", 10, "center", False, False, None),
                ("الكمية التقديرية", "estimated_quantity", 15, "right", False, False, '#,##0.00;(#,##0.00);"-"'),
                ("العلامة التجارية / الموديل المعروض", "supplier_brand", 25, "left", True, False, None),
                ("سعر الوحدة للمورد", "supplier_unit_rate", 15, "right", True, False, '#,##0.00;(#,##0.00);"-"'),
                ("إجمالي السعر للمورد", "supplier_total_price", 18, "right", False, True, '#,##0.00;(#,##0.00);"-"'),
                ("ملاحظات المورد / مدة التوريد", "supplier_remarks", 25, "left", True, False, None),
                ("أساس الكمية", "quantity_basis", 20, "left", False, False, None),
                ("بنود جدول الكميات المرتبطة", "related_boq_items", 20, "center", False, False, None),
                ("الأقسام المصدرية", "source_sections", 15, "left", False, False, None),
                ("مستوى الثقة", "confidence", 12, "center", False, False, None),
                ("ملاحظات", "notes", 30, "left", False, False, None)
            ]
        elif language == "english":
            cols = [
                ("Material Name", "material_name", 25, "left", False, False, None),
                ("Specification", "specification", 40, "left", False, False, None),
                ("Unit", "unit", 10, "center", False, False, None),
                ("Estimated Quantity", "estimated_quantity", 15, "right", False, False, '#,##0.00;(#,##0.00);"-"'),
                ("Supplier Brand / Model", "supplier_brand", 25, "left", True, False, None),
                ("Supplier Unit Rate", "supplier_unit_rate", 15, "right", True, False, '#,##0.00;(#,##0.00);"-"'),
                ("Supplier Total Price", "supplier_total_price", 18, "right", False, True, '#,##0.00;(#,##0.00);"-"'),
                ("Supplier Remarks / Lead Time", "supplier_remarks", 25, "left", True, False, None),
                ("Quantity Basis", "quantity_basis", 20, "left", False, False, None),
                ("Related BOQ Items", "related_boq_items", 20, "center", False, False, None),
                ("Source Sections", "source_sections", 15, "left", False, False, None),
                ("Confidence", "confidence", 12, "center", False, False, None),
                ("Notes", "notes", 30, "left", False, False, None)
            ]
        else: # bilingual
            cols = [
                ("Material Name (EN) / اسم المادة", "material_name_en", 25, "left", False, False, None),
                ("Material Name (AR) / اسم المادة ع", "material_name_ar", 25, "right", False, False, None),
                ("Specification (EN) / المواصفات", "specification_en", 35, "left", False, False, None),
                ("Specification (AR) / المواصفات ع", "specification_ar", 35, "right", False, False, None),
                ("Unit (EN) / الوحدة", "unit_en", 10, "center", False, False, None),
                ("Unit (AR) / الوحدة ع", "unit_ar", 10, "center", False, False, None),
                ("Estimated Qty / الكمية", "estimated_quantity", 15, "right", False, False, '#,##0.00;(#,##0.00);"-"'),
                ("Supplier Brand / Model / العلامة التجارية", "supplier_brand", 25, "left", True, False, None),
                ("Supplier Unit Rate / سعر الوحدة", "supplier_unit_rate", 15, "right", True, False, '#,##0.00;(#,##0.00);"-"'),
                ("Supplier Total Price / إجمالي السعر", "supplier_total_price", 18, "right", False, True, '#,##0.00;(#,##0.00);"-"'),
                ("Supplier Remarks / ملاحظات المورد", "supplier_remarks", 25, "left", True, False, None),
                ("Qty Basis / أساس الكمية", "quantity_basis", 20, "left", False, False, None),
                ("BOQ Refs / البنود", "related_boq_items", 15, "center", False, False, None),
                ("Sources / المصدر", "source_sections", 15, "left", False, False, None),
                ("Confidence / الثقة", "confidence", 12, "center", False, False, None),
                ("Notes / ملاحظات", "notes", 25, "left", False, False, None)
            ]
            
        # Create worksheets for each package
        package_sheet_map = {}
        used_titles = set([ws_summary.title.lower()])
        
        for pkg in sorted(grouped_materials.keys()):
            sanitized_title = ExcelExporter.sanitize_sheet_title(pkg)
            base = sanitized_title
            counter = 1
            while sanitized_title.lower() in used_titles:
                suffix = f"_{counter}"
                sanitized_title = base[:30-len(suffix)] + suffix
                counter += 1
            used_titles.add(sanitized_title.lower())
            package_sheet_map[pkg] = sanitized_title
            
            ws = wb.create_sheet(title=sanitized_title)
            ws.views.sheetView[0].showGridLines = True
            
            if language == "arabic":
                ws.views.sheetView[0].rightToLeft = True
            else:
                ws.views.sheetView[0].rightToLeft = False
                
            # Write headers
            qty_col_letter = None
            rate_col_letter = None
            total_col_letter = None
            
            for col_idx, col_info in enumerate(cols):
                col_num = col_idx + 1
                field_name = col_info[1]
                if field_name == "estimated_quantity":
                    qty_col_letter = get_column_letter(col_num)
                elif field_name == "supplier_unit_rate":
                    rate_col_letter = get_column_letter(col_num)
                elif field_name == "supplier_total_price":
                    total_col_letter = get_column_letter(col_num)
                    
                cell = ws.cell(row=1, column=col_num, value=col_info[0])
                cell.fill = RUST_HEADER_FILL
                cell.font = WHITE_FONT
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = CELL_BORDER
                
            # Write package materials data
            pkg_materials = grouped_materials[pkg]
            for row_idx, mat in enumerate(pkg_materials):
                r_num = row_idx + 2
                row_fill = ZEBRA_FILL if row_idx % 2 == 1 else WHITE_FILL
                
                for col_idx, col_info in enumerate(cols):
                    col_num = col_idx + 1
                    field_name = col_info[1]
                    cell = ws.cell(row=r_num, column=col_num)
                    cell.fill = row_fill
                    cell.border = CELL_BORDER
                    
                    align_h = col_info[3]
                    cell.alignment = Alignment(horizontal=align_h, vertical="center")
                    
                    is_input = col_info[4]
                    is_formula = col_info[5]
                    num_format = col_info[6]
                    
                    if is_formula:
                        if field_name == "supplier_total_price":
                            cell.value = f'=IF(ISBLANK({rate_col_letter}{r_num}), "", {qty_col_letter}{r_num}*{rate_col_letter}{r_num})'
                        cell.font = REGULAR_FONT
                    elif is_input:
                        cell.value = "" # Blank for suppliers to type
                        cell.font = BLUE_FONT # Supplier input font (blue)
                    else:
                        val = mat.get(field_name, "")
                        if isinstance(val, list):
                            val = ", ".join(str(v) for v in val)
                        elif val is None:
                            val = ""
                        
                        if field_name == "estimated_quantity" and val != "":
                            try:
                                cell.value = float(val)
                            except ValueError:
                                cell.value = val
                        else:
                            cell.value = val
                        cell.font = REGULAR_FONT
                        
                    if num_format and cell.value != "":
                        cell.number_format = num_format
            
            # Set column widths & row heights
            for col_idx, col_info in enumerate(cols):
                ws.column_dimensions[get_column_letter(col_idx + 1)].width = col_info[2]
            ws.row_dimensions[1].height = 28
            
        # 2. Write Summary Page Title & Info
        ws_summary.cell(row=1, column=1, value="Project Materials Estimation Summary / ملخص تقدير مواد المشروع").font = Font(name="Segoe UI", size=14, bold=True, color="D35400")
        ws_summary.row_dimensions[1].height = 24
        
        ws_summary.cell(row=3, column=1, value="Source File / الملف المصدر:").font = BOLD_FONT
        ws_summary.cell(row=3, column=2, value=os.path.basename(output_path).replace(".xlsx", "")).font = REGULAR_FONT
        
        ws_summary.cell(row=4, column=1, value="Date Extracted / تاريخ الاستخراج:").font = BOLD_FONT
        ws_summary.cell(row=4, column=2, value=datetime.datetime.now().strftime("%Y-%m-%d %H:%M")).font = REGULAR_FONT
        
        # Package worksheets index table
        ws_summary.cell(row=6, column=1, value="Package Category / باقة المواد").fill = RUST_HEADER_FILL
        ws_summary.cell(row=6, column=1).font = WHITE_FONT
        ws_summary.cell(row=6, column=1).alignment = Alignment(horizontal="center")
        ws_summary.cell(row=6, column=1).border = CELL_BORDER
        
        ws_summary.cell(row=6, column=2, value="Item Count / عدد المواد").fill = RUST_HEADER_FILL
        ws_summary.cell(row=6, column=2).font = WHITE_FONT
        ws_summary.cell(row=6, column=2).alignment = Alignment(horizontal="center")
        ws_summary.cell(row=6, column=2).border = CELL_BORDER
        
        ws_summary.cell(row=6, column=3, value="Link / الرابط").fill = RUST_HEADER_FILL
        ws_summary.cell(row=6, column=3).font = WHITE_FONT
        ws_summary.cell(row=6, column=3).alignment = Alignment(horizontal="center")
        ws_summary.cell(row=6, column=3).border = CELL_BORDER
        ws_summary.row_dimensions[6].height = 22
        
        # Write packages index rows
        start_row = 7
        for idx, (pkg_name, sheet_title) in enumerate(sorted(package_sheet_map.items())):
            r_num = start_row + idx
            row_fill = ZEBRA_FILL if idx % 2 == 1 else WHITE_FILL
            
            c_name = ws_summary.cell(row=r_num, column=1, value=pkg_name)
            c_name.fill = row_fill
            c_name.font = BOLD_FONT
            c_name.border = CELL_BORDER
            c_name.alignment = Alignment(horizontal="left" if language != "arabic" else "right")
            
            # Count formula referencing Col A of package sheet
            c_count = ws_summary.cell(row=r_num, column=2, value=f"=COUNTA('{sheet_title}'!A:A)-1")
            c_count.fill = row_fill
            c_count.font = GREEN_FONT
            c_count.border = CELL_BORDER
            c_count.alignment = Alignment(horizontal="center")
            
            link_formula = f'=HYPERLINK("#\'{sheet_title}\'!A1", "Go to Sheet / اذهب للصفحة")'
            c_link = ws_summary.cell(row=r_num, column=3, value=link_formula)
            c_link.fill = row_fill
            c_link.font = GREEN_FONT
            c_link.border = CELL_BORDER
            c_link.alignment = Alignment(horizontal="center")
            
        # Total Summary Row
        tot_row = start_row + len(package_sheet_map)
        ws_summary.cell(row=tot_row, column=1, value="Total Extracted Materials / إجمالي المواد المستخرجة").font = BOLD_FONT
        ws_summary.cell(row=tot_row, column=1).border = CELL_BORDER
        ws_summary.cell(row=tot_row, column=1).fill = ZEBRA_FILL
        ws_summary.cell(row=tot_row, column=1).alignment = Alignment(horizontal="left" if language != "arabic" else "right")
        
        sum_formula = f"=SUM(B7:B{tot_row-1})" if len(package_sheet_map) > 0 else "0"
        c_sum = ws_summary.cell(row=tot_row, column=2, value=sum_formula)
        c_sum.font = Font(name="Segoe UI", size=10, bold=True)
        c_sum.border = CELL_BORDER
        c_sum.fill = ZEBRA_FILL
        c_sum.alignment = Alignment(horizontal="center")
        
        c_empty = ws_summary.cell(row=tot_row, column=3, value="")
        c_empty.border = CELL_BORDER
        c_empty.fill = ZEBRA_FILL
        
        ws_summary.column_dimensions['A'].width = 40
        ws_summary.column_dimensions['B'].width = 25
        ws_summary.column_dimensions['C'].width = 25
        
        # 3. Review Flags Sheet
        ws_flags = wb.create_sheet(title="Review Flags" if language != "arabic" else "مؤشرات المراجعة")
        ws_flags.views.sheetView[0].showGridLines = True
        if language == "arabic":
            ws_flags.views.sheetView[0].rightToLeft = True
            
        flag_cols = [
            ("Material Name / اسم المادة", "material_name", 30),
            ("Issue Type / نوع المشكلة", "issue_type", 20),
            ("Description / التفاصيل", "description", 50),
            ("Severity / الأهمية", "severity", 12)
        ]
        
        for c_idx, col_info in enumerate(flag_cols):
            cell = ws_flags.cell(row=1, column=c_idx + 1, value=col_info[0])
            cell.fill = RUST_HEADER_FILL
            cell.font = WHITE_FONT
            cell.border = CELL_BORDER
            cell.alignment = Alignment(horizontal="center")
            
        review_flags = data.get("review_flags", [])
        for r_idx, flag in enumerate(review_flags):
            r_num = r_idx + 2
            row_fill = ZEBRA_FILL if r_idx % 2 == 1 else WHITE_FILL
            for c_idx, col_info in enumerate(flag_cols):
                val = flag.get(col_info[1], "")
                cell = ws_flags.cell(row=r_num, column=c_idx + 1, value=val)
                cell.fill = row_fill
                cell.font = REGULAR_FONT
                cell.border = CELL_BORDER
                cell.alignment = Alignment(horizontal="left" if col_info[1] != "severity" else "center")
        
        for c_idx, col_info in enumerate(flag_cols):
            ws_flags.column_dimensions[get_column_letter(c_idx + 1)].width = col_info[2]
            
        # 4. Assumptions Sheet
        ws_assump = wb.create_sheet(title="Assumptions" if language != "arabic" else "الافتراضات")
        ws_assump.views.sheetView[0].showGridLines = True
        if language == "arabic":
            ws_assump.views.sheetView[0].rightToLeft = True
            
        assump_cols = [
            ("Scope/Material / المادة أو النطاق", "scope_or_material", 30),
            ("Assumption Made / الافتراض الموضوع", "assumption_made", 50),
            ("Justification / المبرر", "justification", 40)
        ]
        
        for c_idx, col_info in enumerate(assump_cols):
            cell = ws_assump.cell(row=1, column=c_idx + 1, value=col_info[0])
            cell.fill = RUST_HEADER_FILL
            cell.font = WHITE_FONT
            cell.border = CELL_BORDER
            cell.alignment = Alignment(horizontal="center")
            
        assumptions = data.get("assumptions", [])
        for r_idx, ass in enumerate(assumptions):
            r_num = r_idx + 2
            row_fill = ZEBRA_FILL if r_idx % 2 == 1 else WHITE_FILL
            for c_idx, col_info in enumerate(assump_cols):
                val = ass.get(col_info[1], "")
                cell = ws_assump.cell(row=r_num, column=c_idx + 1, value=val)
                cell.fill = row_fill
                cell.font = REGULAR_FONT
                cell.border = CELL_BORDER
                cell.alignment = Alignment(horizontal="left")
                
        for c_idx, col_info in enumerate(assump_cols):
            ws_assump.column_dimensions[get_column_letter(c_idx + 1)].width = col_info[2]
            
        # 5. Warnings Sheet
        ws_warn = wb.create_sheet(title="Warnings" if language != "arabic" else "التحذيرات")
        ws_warn.views.sheetView[0].showGridLines = True
        if language == "arabic":
            ws_warn.views.sheetView[0].rightToLeft = True
            
        cell = ws_warn.cell(row=1, column=1, value="Warning Details / تفاصيل التحذير")
        cell.fill = RUST_HEADER_FILL
        cell.font = WHITE_FONT
        cell.border = CELL_BORDER
        cell.alignment = Alignment(horizontal="center")
        
        warnings = data.get("warnings", [])
        for r_idx, warn in enumerate(warnings):
            r_num = r_idx + 2
            row_fill = ZEBRA_FILL if r_idx % 2 == 1 else WHITE_FILL
            cell = ws_warn.cell(row=r_num, column=1, value=warn)
            cell.fill = row_fill
            cell.font = REGULAR_FONT
            cell.border = CELL_BORDER
            cell.alignment = Alignment(horizontal="left")
            
        ws_warn.column_dimensions['A'].width = 100
        
        # Save workbook
        wb.save(output_path)
        logger.info(f"Excel successfully saved to {output_path}")

# CSV Exporter for ERP systems
class CSVExporter:
    @staticmethod
    def export(data: Dict[str, Any], output_path: str, language: str) -> None:
        """Generates a flat, unstyled CSV file of all materials for ERP system ingestion."""
        import csv
        materials = data.get("materials", [])
        
        # Decide headers and fields based on language selection
        if language == "arabic":
            headers = ["الباقة", "اسم المادة", "المواصفات الفنية", "الوحدة", "الكمية التقديرية", "أساس الكمية", "البنود المرتبطة", "الأقسام المصدرية", "مستوى الثقة", "ملاحظات"]
            fields = ["package", "material_name", "specification", "unit", "estimated_quantity", "quantity_basis", "related_boq_items", "source_sections", "confidence", "notes"]
        elif language == "english":
            headers = ["Package", "Material Name", "Specification", "Unit", "Estimated Quantity", "Quantity Basis", "Related BOQ Items", "Source Sections", "Confidence", "Notes"]
            fields = ["package", "material_name", "specification", "unit", "estimated_quantity", "quantity_basis", "related_boq_items", "source_sections", "confidence", "notes"]
        else: # bilingual
            headers = [
                "Package (EN)", "Package (AR)",
                "Material Name (EN)", "Material Name (AR)",
                "Specification (EN)", "Specification (AR)",
                "Unit (EN)", "Unit (AR)",
                "Estimated Quantity", "Quantity Basis",
                "BOQ Refs", "Sources", "Confidence", "Notes"
            ]
            fields = [
                "package_en", "package_ar",
                "material_name_en", "material_name_ar",
                "specification_en", "specification_ar",
                "unit_en", "unit_ar",
                "estimated_quantity", "quantity_basis",
                "related_boq_items", "source_sections", "confidence", "notes"
            ]
            
        try:
            with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow(headers)
                
                for mat in materials:
                    row = []
                    for field in fields:
                        val = mat.get(field, "")
                        if val is None:
                            val = ""
                        elif isinstance(val, list):
                            val = ", ".join(str(v) for v in val)
                        row.append(val)
                    writer.writerow(row)
            logger.info(f"CSV successfully saved to {output_path}")
        except Exception as e:
            logger.error(f"Failed to generate CSV file: {e}")
            raise e

# PDF Exporter using Edge Headless Printing
class PDFExporter:
    @staticmethod
    def export(data: Dict[str, Any], output_path: str, language: str, config: Dict[str, Any]) -> None:
        """Generates a print-optimized PDF from the structured materials list using headless Edge."""
        import subprocess
        import tempfile
        
        materials = data.get("materials", [])
        if not materials:
            logger.warning("No materials extracted. Creating a blank PDF report.")
            
        # Define columns and labels
        if language == "arabic":
            cols = [
                ("material_name", "اسم المادة"),
                ("specification", "المواصفات الفنية"),
                ("unit", "الوحدة"),
                ("estimated_quantity", "الكمية"),
                ("quantity_basis", "أساس الكمية"),
                ("notes", "ملاحظات")
            ]
            dir_attr = "rtl"
        elif language == "english":
            cols = [
                ("material_name", "Material Name"),
                ("specification", "Specification"),
                ("unit", "Unit"),
                ("estimated_quantity", "Qty"),
                ("quantity_basis", "Qty Basis"),
                ("notes", "Notes")
            ]
            dir_attr = "ltr"
        else: # bilingual
            cols = [
                ("material_name_en", "Material (EN)"),
                ("material_name_ar", "اسم المادة"),
                ("specification_en", "Specification (EN)"),
                ("specification_ar", "المواصفات ع"),
                ("unit_en", "Unit"),
                ("estimated_quantity", "Qty"),
                ("notes", "Notes")
            ]
            dir_attr = "ltr"

        # Separate keys and get print settings
        col_keys = [c[0] for c in cols]
        settings = AIService.get_print_settings(config, col_keys, materials)
        
        # Build grouped materials by package for layout sectioning
        grouped_materials = {}
        for mat in materials:
            if language == "arabic":
                pkg = mat.get("package") or "عام"
            elif language == "english":
                pkg = mat.get("package") or "General"
            else: # bilingual
                pkg = mat.get("package_en") or mat.get("package") or "General"
            pkg = str(pkg).strip()
            if not pkg:
                pkg = "General" if language != "arabic" else "عام"
            if pkg not in grouped_materials:
                grouped_materials[pkg] = []
            grouped_materials[pkg].append(mat)
            
        # Compile HTML content
        html_parts = []
        html_parts.append(f"""<!DOCTYPE html>
<html dir="{dir_attr}" lang="{ 'ar' if language == 'arabic' else 'en' }">
<head>
    <meta charset="utf-8">
    <title>Materials Estimation Report</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&family=Cairo:wght@300;400;600;700&display=swap" rel="stylesheet">
    <style>
        @page {{
            size: letter {settings.get('orientation', 'landscape')};
            margin: 1.2cm;
        }}
        body {{
            font-family: 'Outfit', 'Cairo', Arial, sans-serif;
            margin: 0;
            padding: 0;
            font-size: {settings.get('font_size', '11px')};
            color: #2C3E50;
            background: #fff;
            line-height: 1.4;
        }}
        .report-header {{
            display: flex;
            justify-content: space-between;
            align-items: center;
            border-bottom: 2px solid #D35400;
            padding-bottom: 12px;
            margin-bottom: 20px;
        }}
        .report-title {{
            font-size: 20px;
            font-weight: 700;
            color: #D35400;
            margin: 0;
        }}
        .report-metadata {{
            text-align: { 'left' if language == 'arabic' else 'right' };
            font-size: 10px;
            color: #7F8C8D;
        }}
        .package-section {{
            page-break-after: always;
        }}
        .package-section:last-of-type {{
            page-break-after: avoid;
        }}
        .package-title {{
            font-size: 14px;
            font-weight: 600;
            color: #2C3E50;
            margin-top: 15px;
            margin-bottom: 8px;
            border-bottom: 1px solid #BDC3C7;
            padding-bottom: 4px;
            text-transform: uppercase;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin-bottom: 15px;
            page-break-inside: auto;
        }}
        tr {{
            page-break-inside: avoid;
            page-break-after: auto;
        }}
        thead {{
            display: table-header-group;
        }}
        th {{
            background-color: #D35400;
            color: #ffffff;
            font-weight: 600;
            text-align: center;
            padding: 6px 8px;
            border: 1px solid #D5D8DC;
            font-size: 10px;
        }}
        td {{
            padding: 5px 8px;
            border: 1px solid #D5D8DC;
            vertical-align: top;
            font-size: {settings.get('font_size', '11px')};
        }}
        tr:nth-child(even) {{
            background-color: #F9EBEA;
        }}
        .text-left {{ text-align: left; }}
        .text-center {{ text-align: center; }}
        .text-right {{ text-align: right; }}
        
        /* Column Widths */
""")
        
        # Inject column widths
        col_widths = settings.get("column_widths", {})
        for col_key, width in col_widths.items():
            html_parts.append(f"        .col-{col_key} {{ width: {width}; }}\n")
            
        # Inject wrap columns
        wrap_cols = settings.get("wrap_columns", [])
        for col_key in wrap_cols:
            html_parts.append(f"        .col-{col_key} {{ white-space: normal; word-break: break-word; }}\n")
            
        html_parts.append("""    </style>
</head>
<body>
""")
        
        # Write each package section
        source_name = os.path.basename(output_path).replace(".pdf", "")
        run_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        
        for pkg, pkg_mats in sorted(grouped_materials.items()):
            html_parts.append(f"""
    <div class="package-section">
        <div class="report-header">
            <div>
                <h1 class="report-title">TAWREED - Quantity Takeoff Report</h1>
                <div style="font-size: 12px; color: #7F8C8D; margin-top: 2px;">{pkg}</div>
            </div>
            <div class="report-metadata">
                <div>File: {source_name}</div>
                <div>Date: {run_date}</div>
            </div>
        </div>
        
        <h2 class="package-title">{pkg}</h2>
        <table>
            <thead>
                <tr>
""")
            # Write Table Headers
            for col_key, label in cols:
                html_parts.append(f'                    <th class="col-{col_key}">{label}</th>\n')
                
            html_parts.append("""                </tr>
            </thead>
            <tbody>
""")
            # Write Table Rows
            for mat in pkg_mats:
                html_parts.append("                <tr>\n")
                for col_key, label in cols:
                    val = mat.get(col_key, "")
                    if val is None:
                        val = ""
                    elif isinstance(val, list):
                        val = ", ".join(str(v) for v in val)
                        
                    # Alignment CSS class
                    align_cls = "text-left"
                    if col_key == "estimated_quantity":
                        align_cls = "text-right"
                        try:
                            # Format estimated quantity neatly
                            val = f"{float(val):,.2f}"
                        except ValueError:
                            pass
                    elif col_key in ["unit", "unit_en", "unit_ar", "confidence"]:
                        align_cls = "text-center"
                    elif language == "arabic" and col_key in ["material_name", "specification"]:
                        align_cls = "text-right"
                    elif language == "bilingual" and col_key in ["material_name_ar", "specification_ar"]:
                        align_cls = "text-right"
                        
                    html_parts.append(f'                    <td class="col-{col_key} {align_cls}">{val}</td>\n')
                html_parts.append("                </tr>\n")
                
            html_parts.append("""            </tbody>
        </table>
    </div>
""")
            
        html_parts.append("""</body>
</html>
""")
        
        # Combine all HTML
        html_content = "".join(html_parts)
        
        # Locate Microsoft Edge
        edge_paths = [
            r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
            r"C:\Program Files\Microsoft\Edge\Application\msedge.exe",
            os.path.expandvars(r"%LocalAppData%\Microsoft\Edge\Application\msedge.exe")
        ]
        edge_exec = None
        for path in edge_paths:
            if os.path.exists(path):
                edge_exec = path
                break
                
        if not edge_exec:
            try:
                out = subprocess.check_output(["where", "msedge"], text=True)
                lines = out.strip().split("\n")
                if lines and os.path.exists(lines[0].strip()):
                    edge_exec = lines[0].strip()
            except Exception:
                pass
                
        if not edge_exec:
            try:
                import winreg
                key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\msedge.exe")
                val, _ = winreg.QueryValue(key, None)
                if val and os.path.exists(val):
                    edge_exec = val
            except Exception:
                pass
                
        if not edge_exec:
            logger.error("Microsoft Edge was not found on this system. Cannot convert HTML to PDF.")
            raise Exception("Microsoft Edge is required to generate PDF reports but was not found on the system.")
            
        # Write to temporary file and execute print
        temp_fd, temp_path = tempfile.mkstemp(suffix=".html", prefix="tawreed_pdf_")
        try:
            with os.fdopen(temp_fd, "w", encoding="utf-8") as tf:
                tf.write(html_content)
                
            # Invoke Edge headless printing
            cmd = [
                edge_exec,
                "--headless",
                "--disable-gpu",
                f"--print-to-pdf={output_path}",
                temp_path
            ]
            logger.info(f"Running Edge Headless PDF Generation command: {cmd}")
            subprocess.run(cmd, check=True, timeout=30)
            logger.info(f"PDF successfully generated at {output_path}")
        except Exception as e:
            logger.error(f"Failed to generate PDF via Edge command: {e}")
            raise e
        finally:
            if os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except Exception as clean_e:
                    logger.warning(f"Failed to delete temp HTML file {temp_path}: {clean_e}")

# The Main Processing Orchestrator
class JobProcessor:
    @staticmethod
    def run(file_path: str, output_dir: str, config: Dict[str, Any], selected_lang: str) -> Dict[str, Any]:
        """Orchestrates the parsing, LLM execution, repair, and excel generation."""
        job_id = str(uuid.uuid4())
        job_log_dir = os.path.join(LOG_DIR, job_id)
        os.makedirs(job_log_dir, exist_ok=True)
        
        logger.info(f"Starting Job {job_id} for file: {file_path}")
        warnings = []
        
        # Step 1: Parse and serialize the document
        file_ext = os.path.splitext(file_path)[1].lower()
        serialized_text = ""
        image_pages = []
        
        try:
            if file_ext in [".xlsx", ".xls", ".xlsm"]:
                serialized_text, parse_warns = DocumentParser.parse_excel(file_path)
                warnings.extend(parse_warns)
            elif file_ext == ".csv":
                serialized_text, parse_warns = DocumentParser.parse_csv(file_path)
                warnings.extend(parse_warns)
            elif file_ext in [".docx", ".doc"]:
                serialized_text, parse_warns = DocumentParser.parse_word(file_path)
                warnings.extend(parse_warns)
            elif file_ext == ".pdf":
                serialized_text, parse_warns, is_scanned = DocumentParser.parse_digital_pdf(file_path)
                warnings.extend(parse_warns)
                if is_scanned:
                    logger.info("PDF is scanned. Rendering pages to images for AI Vision OCR...")
                    warnings.append("This is a scanned PDF. Pages were rendered to images and sent to the AI Vision model.")
                    image_pages = DocumentParser.render_pdf_pages_to_images(file_path)
            elif file_ext in [".png", ".jpg", ".jpeg"]:
                logger.info("Direct image upload. Reading bytes for AI Vision...")
                with open(file_path, "rb") as f:
                    img_bytes = f.read()
                image_pages = [(1, img_bytes)]
                serialized_text = f"Source Image File: {os.path.basename(file_path)}"
            else: # Text/fallback
                serialized_text, parse_warns = DocumentParser.parse_text(file_path)
                warnings.extend(parse_warns)
        except Exception as e:
            error_msg = f"Failed to parse document: {str(e)}"
            logger.error(error_msg)
            raise Exception(error_msg)

        # Check content limits
        max_chars = config.get("max_content_length_chars", 80000)
        if len(serialized_text) > max_chars:
            logger.warning(f"Content length ({len(serialized_text)}) exceeds configured limit ({max_chars}). Truncating.")
            serialized_text = serialized_text[:max_chars] + "\n\n... [CONTENT TRUNCATED DUE TO SIZE LIMIT] ..."
            warnings.append("The serialized file content exceeded the character limit and was truncated before sending to the AI.")

        # Save serialized content for diagnostics
        with open(os.path.join(job_log_dir, "serialized_input.txt"), "w", encoding="utf-8") as f:
            f.write(serialized_text)
            
        # Step 2: Construct prompt and call AI
        prompt = AIService.get_system_prompt(selected_lang)
        
        # Save prompt for diagnostics
        with open(os.path.join(job_log_dir, "system_prompt.txt"), "w", encoding="utf-8") as f:
            f.write(prompt)
            
        image_bytes_only = [p[1] for p in image_pages] if image_pages else None
        
        # API call
        raw_ai_res, error_msg = AIService.call_ai(config, prompt, text_content=serialized_text, image_bytes_list=image_bytes_only)
        
        # Save raw response for diagnostics
        with open(os.path.join(job_log_dir, "raw_ai_response.txt"), "w", encoding="utf-8", errors="ignore") as f:
            f.write(raw_ai_res or "[No Response]")
        if error_msg:
            with open(os.path.join(job_log_dir, "error.log"), "w", encoding="utf-8") as f:
                f.write(error_msg)
            raise Exception(f"AI Provider returned an error: {error_msg}")

        # Step 3: Parse and Repair JSON
        parsed_data = None
        json_error = ""
        try:
            parsed_data = JSONRepairService.repair_json(raw_ai_res)
        except Exception as je:
            json_error = str(je)
            logger.error(f"JSON validation failed: {json_error}. Attempting automated AI repair query...")
            
            # Save repair attempt log
            with open(os.path.join(job_log_dir, "repair_attempt.log"), "w", encoding="utf-8") as f:
                f.write(f"Initial JSON parse error: {json_error}\n")
            
            # Send repair request
            try:
                parsed_data = JSONRepairService.attempt_ai_repair(config, raw_ai_res, json_error)
                logger.info("AI successfully repaired the JSON response.")
                with open(os.path.join(job_log_dir, "repair_attempt.log"), "a", encoding="utf-8") as f:
                    f.write("AI repair SUCCESSFUL.\n")
                warnings.append("The initial AI response had formatting issues and was automatically repaired.")
            except Exception as repair_e:
                repair_err_msg = f"Failed to repair AI response JSON: {str(repair_e)}"
                logger.error(repair_err_msg)
                with open(os.path.join(job_log_dir, "repair_attempt.log"), "a", encoding="utf-8") as f:
                    f.write(f"AI repair FAILED. Error: {repair_err_msg}\n")
                raise Exception("The AI returned a malformed response that could not be parsed or repaired.")

        # Incorporate backend warnings into parsed data warnings
        if "warnings" not in parsed_data or not isinstance(parsed_data["warnings"], list):
            parsed_data["warnings"] = []
        parsed_data["warnings"].extend(warnings)

        # Save clean JSON output for diagnostics
        with open(os.path.join(job_log_dir, "extracted_data.json"), "w", encoding="utf-8") as f:
            json.dump(parsed_data, f, indent=4, ensure_ascii=False)

        # Step 4: Export to Excel and CSV
        base_name_no_ext = os.path.splitext(os.path.basename(file_path))[0]
        
        out_base_name = f"Tawreed_{base_name_no_ext}.xlsx"
        output_path = os.path.join(output_dir, out_base_name)
        
        out_csv_name = f"Tawreed_{base_name_no_ext}_for_erp.csv"
        output_csv = os.path.join(output_dir, out_csv_name)
        
        # Export XLSX
        try:
            ExcelExporter.export(parsed_data, output_path, selected_lang)
        except Exception as ee:
            logger.error(f"Excel export failed: {ee}\n{traceback.format_exc()}")
            raise Exception(f"Failed to generate Excel file: {str(ee)}")
            
        # Export CSV
        try:
            CSVExporter.export(parsed_data, output_csv, selected_lang)
        except Exception as ce:
            logger.error(f"CSV export failed: {ce}\n{traceback.format_exc()}")
            raise Exception(f"Failed to generate CSV file: {str(ce)}")

        # Step 5: Save job manifest
        manifest = {
            "job_id": job_id,
            "timestamp": datetime.datetime.now().isoformat(),
            "source_file": file_path,
            "output_file": output_path,
            "output_csv": output_csv,
            "output_dir": output_dir,
            "status": "success",
            "summary": parsed_data.get("summary", {}),
            "warnings_count": len(parsed_data.get("warnings", [])),
            "flags_count": len(parsed_data.get("review_flags", []))
        }
        
        with open(os.path.join(job_log_dir, "manifest.json"), "w", encoding="utf-8") as f:
            json.dump(manifest, f, indent=4, ensure_ascii=False)

        return manifest
